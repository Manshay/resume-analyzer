from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Twips
import os
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

class EnhancedResumeGenerator:
    def __init__(self, template_type):
        self.template_type = template_type
        self.document = Document()
        self.style_config = self._get_style_config()
        self._setup_styles()
    
    def _get_style_config(self):
        """Get template-specific style configuration"""
        styles = {
            "Professional": {
                "colors": {"primary": RGBColor(0, 51, 102), "secondary": RGBColor(128, 128, 128)},
                "fonts": {"heading": "Calibri", "body": "Calibri"},
                "spacing": {"before": Pt(12), "after": Pt(8)},
                "margins": {"top": 0.5, "left": 0.7, "bottom": 0.5, "right": 0.7}
            },
            "Modern": {
                "colors": {"primary": RGBColor(51, 51, 51), "secondary": RGBColor(0, 102, 204)},
                "fonts": {"heading": "Arial", "body": "Arial"},
                "spacing": {"before": Pt(14), "after": Pt(10)},
                "margins": {"top": 0.6, "left": 0.75, "bottom": 0.6, "right": 0.75}
            },
            "Creative": {
                "colors": {"primary": RGBColor(76, 175, 80), "secondary": RGBColor(33, 33, 33)},
                "fonts": {"heading": "Georgia", "body": "Helvetica"},
                "spacing": {"before": Pt(14), "after": Pt(10)},
                "margins": {"top": 0.8, "left": 1.0, "bottom": 0.8, "right": 1.0}
            },
            "Academic": {
                "colors": {"primary": RGBColor(140, 20, 20), "secondary": RGBColor(68, 68, 68)},
                "fonts": {"heading": "Times New Roman", "body": "Times New Roman"},
                "spacing": {"before": Pt(12), "after": Pt(6)},
                "margins": {"top": 1.0, "left": 1.2, "bottom": 1.0, "right": 1.2}
            },
            "Technical": {
                "colors": {"primary": RGBColor(25, 118, 210), "secondary": RGBColor(66, 66, 66)},
                "fonts": {"heading": "Consolas", "body": "Segoe UI"},
                "spacing": {"before": Pt(10), "after": Pt(8)},
                "margins": {"top": 0.5, "left": 0.7, "bottom": 0.5, "right": 0.7}
            }
        }
        return styles.get(self.template_type, styles["Professional"])
    
    def _setup_styles(self):
        """Setup all document styles based on template"""
        try:
            config = self.style_config
            
            # Header Style
            header_style = self.document.styles.add_style('CustomHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = config["fonts"]["heading"]
            header_style.font.size = Pt(24)
            header_style.font.color.rgb = config["colors"]["primary"]
            header_style.font.bold = True
            header_style.paragraph_format.space_after = Pt(4)
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Heading Style (Section titles)
            heading_style = self.document.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = config["fonts"]["heading"]
            heading_style.font.size = Pt(16)
            heading_style.font.color.rgb = config["colors"]["primary"]
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = config["spacing"]["before"]
            heading_style.paragraph_format.space_after = config["spacing"]["after"]
            
            # Subheading Style (Job titles, school names)
            subheading_style = self.document.styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
            subheading_style.font.name = config["fonts"]["body"]
            subheading_style.font.size = Pt(12)
            subheading_style.font.bold = True
            subheading_style.font.color.rgb = config["colors"]["secondary"]
            subheading_style.paragraph_format.space_before = Pt(6)
            subheading_style.paragraph_format.space_after = Pt(2)
            
            # Body Style
            body_style = self.document.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = config["fonts"]["body"]
            body_style.font.size = Pt(11)
            body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            body_style.paragraph_format.space_after = Pt(3)
            
            # Set document margins
            sections = self.document.sections
            for section in sections:
                section.top_margin = Inches(config["margins"]["top"])
                section.bottom_margin = Inches(config["margins"]["bottom"])
                section.left_margin = Inches(config["margins"]["left"])
                section.right_margin = Inches(config["margins"]["right"])
                
        except Exception as e:
            logger.error(f"Error setting up styles: {e}")

    def _add_horizontal_line(self):
        """Add a decorative horizontal line"""
        try:
            paragraph = self.document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(12)
            run = paragraph.add_run('_' * 100)
            run.font.color.rgb = RGBColor(200, 200, 200)
            run.font.size = Pt(8)
        except Exception as e:
            logger.error(f"Error adding horizontal line: {e}")

    def add_header(self, name, contact_info):
        """Add header based on template type"""
        try:
            if self.template_type == "Creative":
                self._add_creative_header(name, contact_info)
            elif self.template_type == "Technical":
                self._add_technical_header(name, contact_info)
            else:
                self._add_standard_header(name, contact_info)
        except Exception as e:
            logger.error(f"Error adding header: {e}")

    def _add_standard_header(self, name, contact_info):
        """Standard header for Professional, Modern, Academic templates"""
        header = self.document.add_paragraph(style='CustomHeader')
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = header.add_run(name)
        name_run.bold = True
        
        self._add_horizontal_line()
        
        contact_paragraph = self.document.add_paragraph(style='CustomBody')
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = []
        
        if contact_info.get('email'):
            contact_text.append(contact_info['email'])
        if contact_info.get('phone'):
            contact_text.append(contact_info['phone'])
        if contact_info.get('location'):
            contact_text.append(contact_info['location'])
        if contact_info.get('linkedin'):
            contact_text.append(contact_info['linkedin'])
        if contact_info.get('portfolio'):
            contact_text.append(contact_info['portfolio'])
        
        contact_paragraph.add_run(' | '.join(contact_text))
        self.document.add_paragraph()  # Spacing

    def _add_creative_header(self, name, contact_info):
        """Creative header with table layout"""
        header = self.document.add_paragraph(style='CustomHeader')
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = header.add_run(name.upper())
        name_run.bold = True
        
        # Create contact table
        table = self.document.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = True
        
        try:
            # Left column
            cell = table.cell(0, 0)
            p = cell.paragraphs[0]
            p.add_run(f"📧 {contact_info.get('email', 'N/A')}\n")
            p.add_run(f"📱 {contact_info.get('phone', 'N/A')}")
            
            # Right column
            cell = table.cell(0, 1)
            p = cell.paragraphs[0]
            p.add_run(f"📍 {contact_info.get('location', 'N/A')}\n")
            if contact_info.get('linkedin'):
                p.add_run(f"🔗 {contact_info['linkedin']}")
        except Exception as e:
            logger.error(f"Error creating creative header table: {e}")

    def _add_technical_header(self, name, contact_info):
        """Technical header with code-style formatting"""
        header = self.document.add_paragraph(style='CustomHeader')
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = header.add_run(f">>> {name}")
        name_run.bold = True
        
        contact = self.document.add_paragraph(style='CustomBody')
        contact.add_run("{\n")
        contact.add_run(f'    "email": "{contact_info.get("email", "N/A")}",\n')
        contact.add_run(f'    "phone": "{contact_info.get("phone", "N/A")}",\n')
        contact.add_run(f'    "location": "{contact_info.get("location", "N/A")}",\n')
        if contact_info.get('linkedin'):
            contact.add_run(f'    "linkedin": "{contact_info["linkedin"]}"\n')
        contact.add_run("}")

    def add_summary(self, summary):
        """Add professional summary section"""
        try:
            if not summary or not summary.strip():
                return
                
            heading = self.document.add_heading('Professional Summary', level=1)
            heading.style = 'CustomHeading'
            
            summary_para = self.document.add_paragraph(summary, style='CustomBody')
            summary_para.paragraph_format.space_after = Pt(12)
        except Exception as e:
            logger.error(f"Error adding summary: {e}")

    def add_experience(self, experiences):
        """Add work experience section"""
        try:
            if not experiences or len(experiences) == 0:
                return
                
            heading = self.document.add_heading('Work Experience', level=1)
            heading.style = 'CustomHeading'
            
            if self.template_type == "Creative":
                self._add_experience_creative(experiences)
            elif self.template_type == "Technical":
                self._add_experience_technical(experiences)
            else:
                self._add_experience_standard(experiences)
                
        except Exception as e:
            logger.error(f"Error adding experience: {e}")

    def _add_experience_standard(self, experiences):
        """Standard experience format"""
        for exp in experiences:
            try:
                # Company and Title
                p = self.document.add_paragraph(style='CustomSubheading')
                p.add_run(f"{exp.get('company', 'N/A')} ").bold = True
                p.add_run("| ").italic = True
                p.add_run(f"{exp.get('title', 'N/A')}").bold = True
                
                # Date
                date_p = self.document.add_paragraph(style='CustomBody')
                date_p.add_run(f"{exp.get('start', 'N/A')} - {exp.get('end', 'N/A')}")
                date_p.paragraph_format.left_indent = Inches(0.25)
                
                # Responsibilities
                resp_p = self.document.add_paragraph(style='CustomBody')
                resp_p.add_run(exp.get('responsibilities', 'N/A'))
                resp_p.paragraph_format.left_indent = Inches(0.25)
                
                self.document.add_paragraph()  # Spacing
            except Exception as e:
                logger.error(f"Error adding standard experience: {e}")

    def _add_experience_creative(self, experiences):
        """Creative experience with table layout"""
        for exp in experiences:
            try:
                table = self.document.add_table(rows=1, cols=2)
                table.allow_autofit = True
                
                # Left column: Dates
                left_cell = table.cell(0, 0)
                left_cell.width = Inches(2)
                p = left_cell.paragraphs[0]
                p.add_run(f"{exp.get('start', 'N/A')} -\n{exp.get('end', 'N/A')}").bold = True
                p.add_run(f"\n{exp.get('company', 'N/A')}")
                
                # Right column: Title and Responsibilities
                right_cell = table.cell(0, 1)
                p = right_cell.paragraphs[0]
                p.add_run(f"{exp.get('title', 'N/A')}\n").bold = True
                p.add_run(exp.get('responsibilities', 'N/A'))
                
                self.document.add_paragraph()
            except Exception as e:
                logger.error(f"Error adding creative experience: {e}")

    def _add_experience_technical(self, experiences):
        """Technical experience with code-style format"""
        for exp in experiences:
            try:
                p = self.document.add_paragraph(style='CustomSubheading')
                company_clean = exp.get('company', 'Company').replace(' ', '')
                p.add_run(f"class {company_clean}_Role {{\n")
                p.add_run(f"    position: {exp.get('title', 'N/A')};\n")
                p.add_run(f"    period: {exp.get('start', 'N/A')} -> {exp.get('end', 'N/A')};\n")
                p.add_run("    responsibilities: {\n")
                
                responsibilities = exp.get('responsibilities', '')
                for resp in responsibilities.split('\n'):
                    if resp.strip():
                        p.add_run(f"        • {resp.strip()}\n")
                p.add_run("    }\n}")
                
                self.document.add_paragraph()
            except Exception as e:
                logger.error(f"Error adding technical experience: {e}")

    def add_education(self, education):
        """Add education section"""
        try:
            if not education or len(education) == 0:
                return
                
            heading = self.document.add_heading('Education', level=1)
            heading.style = 'CustomHeading'
            
            if self.template_type == "Academic":
                self._add_education_academic(education)
            else:
                self._add_education_standard(education)
                
        except Exception as e:
            logger.error(f"Error adding education: {e}")

    def _add_education_standard(self, education):
        """Standard education format"""
        for edu in education:
            try:
                p = self.document.add_paragraph(style='CustomSubheading')
                p.add_run(f"{edu.get('institution', 'N/A')} - {edu.get('degree', 'N/A')}").bold = True
                
                date_p = self.document.add_paragraph(style='CustomBody')
                date_p.add_run(f"{edu.get('start', 'N/A')} - {edu.get('end', 'N/A')}")
                date_p.paragraph_format.left_indent = Inches(0.25)
                
                if edu.get('description'):
                    desc_p = self.document.add_paragraph(edu['description'], style='CustomBody')
                    desc_p.paragraph_format.left_indent = Inches(0.25)
                    
                self.document.add_paragraph()
            except Exception as e:
                logger.error(f"Error adding standard education: {e}")

    def _add_education_academic(self, education):
        """Academic education format with detailed styling"""
        for edu in education:
            try:
                p = self.document.add_paragraph(style='CustomSubheading')
                p.add_run(f"{edu.get('degree', 'N/A')}\n").bold = True
                p.add_run(f"{edu.get('institution', 'N/A')}").italic = True
                
                date_p = self.document.add_paragraph(style='CustomBody')
                date_p.add_run(f"{edu.get('start', 'N/A')} - {edu.get('end', 'N/A')}")
                
                if edu.get('description'):
                    desc_p = self.document.add_paragraph(edu['description'], style='CustomBody')
                    desc_p.paragraph_format.left_indent = Inches(0.25)
                    
                self.document.add_paragraph()
            except Exception as e:
                logger.error(f"Error adding academic education: {e}")

    def add_skills(self, skills):
        """Add skills section"""
        try:
            if not skills or (isinstance(skills, str) and not skills.strip()):
                return
                
            heading = self.document.add_heading('Skills', level=1)
            heading.style = 'CustomHeading'
            
            if isinstance(skills, str):
                skills_list = [s.strip() for s in skills.split(',') if s.strip()]
            else:
                skills_list = skills
            
            if self.template_type == "Technical":
                self._add_skills_technical(skills_list)
            elif self.template_type == "Creative":
                self._add_skills_creative(skills_list)
            else:
                self._add_skills_standard(skills_list)
                
        except Exception as e:
            logger.error(f"Error adding skills: {e}")

    def _add_skills_standard(self, skills_list):
        """Standard skills with bullet points"""
        p = self.document.add_paragraph(style='CustomBody')
        for skill in skills_list:
            p.add_run(f"• {skill}\n")

    def _add_skills_creative(self, skills_list):
        """Creative skills in grid layout"""
        try:
            rows = (len(skills_list) + 1) // 2
            table = self.document.add_table(rows=rows, cols=2)
            table.allow_autofit = True
            
            i = 0
            for row in table.rows:
                for cell in row.cells:
                    if i < len(skills_list):
                        cell.text = f"▪ {skills_list[i]}"
                    i += 1
        except Exception as e:
            logger.error(f"Error adding creative skills: {e}")

    def _add_skills_technical(self, skills_list):
        """Technical skills in code format"""
        p = self.document.add_paragraph(style='CustomBody')
        p.add_run("const skills = {\n")
        for i, skill in enumerate(skills_list):
            comma = ',' if i < len(skills_list) - 1 else ''
            p.add_run(f'    "{skill}"{comma}\n')
        p.add_run("};")

    def generate(self, resume_data):
        """Generate complete resume from data dictionary"""
        try:
            # Add header
            if resume_data.get('personal_info'):
                self.add_header(
                    resume_data['personal_info'].get('name', 'Your Name'),
                    resume_data['personal_info']
                )
            
            # Add summary
            if resume_data.get('summary'):
                self.add_summary(resume_data['summary'])
            
            # Add experience
            if resume_data.get('experience'):
                self.add_experience(resume_data['experience'])
            
            # Add education
            if resume_data.get('education'):
                self.add_education(resume_data['education'])
            
            # Add skills
            if resume_data.get('skills'):
                self.add_skills(resume_data['skills'])
                
            return True
        except Exception as e:
            logger.error(f"Error generating resume: {e}")
            return False

    def save(self, output_path):
        """Save resume in both DOCX and PDF formats"""
        try:
            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
            
            # Save DOCX
            docx_path = f"{output_path}.docx"
            self.document.save(docx_path)
            logger.info(f"Resume saved to {docx_path}")
            
            # Try PDF conversion
            pdf_path = f"{output_path}.pdf"
            try:
                import pdfkit
                options = {
                    'page-size': 'Letter',
                    'margin-top': '0.5in',
                    'margin-right': '0.5in',
                    'margin-bottom': '0.5in',
                    'margin-left': '0.5in',
                    'encoding': 'UTF-8',
                    'quiet': ''
                }
                pdfkit.from_file(docx_path, pdf_path, options=options)
                logger.info(f"Resume converted to PDF: {pdf_path}")
            except Exception as pdf_error:
                logger.warning(f"PDF conversion failed: {pdf_error}. DOCX file saved successfully.")
            
            return {
                'docx_path': docx_path,
                'pdf_path': pdf_path if os.path.exists(pdf_path) else None,
                'success': True
            }
        except Exception as e:
            logger.error(f"Error saving resume: {e}")
            return {
                'success': False,
                'error': str(e)
            }